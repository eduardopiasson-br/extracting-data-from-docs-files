from docx import Document
import os, re, csv

# Função de Leitura do Arquivo e Estruturação do CSV
def percorrer_pasta(caminho):
    database = []
    header = ["CAMINHO", "DATA", "GENERO", "IDADE", "PAÍS", "RELIGIÃO", "RAÇA/COR", "POSSUI DOC.", "TEM FILHOS", "PCD OU ESPECIAL", "PROB. SAÚDE OU DOENÇA", "REALIZANDO TRAT. SAÚDE", "USA MEDICAMENTO", "REALIZOU CIRURGIA", "ACOMP. PSICOLÓGICO", "ATEND. MENTAL/PSIQ", "USA DROGAS", "TRAT. P/ DEP. QUIMICA", "ESCOLARIDADE", "ATV. REMUNERADA", "BOLSA FAMÍLIA", "ID JOVEM", "RECEBE BPC", "BENEFICIO ASS. SOCIAL", "BENEFICIO PREV. PRIVADA",  "EGRESSO SIS. PENAL", "VINCULO FAMILIAR ROMP.", "ABANDONO", "FOI ROUBADO OU FURTADO", "DEPEND. ALCOOL DROGA", "ESTRANGEIRO", "DESEMPREGO", "SAÚDE MENTAL", "TRECHEIRO", "MIGRANTE", "TEMPO DE RUA", "TEMPO DE RUA EM FOZ", "MOTIVO VINDA FOZ", "IR P/ OUTRO CEP", "MORAVA EM"]
    database += [header]

    for root, dirs, files in os.walk(caminho):
        for file in files:
            if file.endswith(".docx") or file.endswith(".doc"):
                caminho_arquivo = os.path.join(root, file)
                textinho = extrair_texto_do_docx(caminho_arquivo)
                new_data = [raspagem_texto(textinho, caminho_arquivo)]
                if new_data != [[]]:
                    print(caminho_arquivo + ' = PROCESSADO')
                    database += new_data

    csv.register_dialect('myDialect', delimiter=';', quoting=csv.QUOTE_ALL)
    with open('registro_pacientes.csv', 'w', newline='') as file:
        writer = csv.writer(file, dialect='myDialect')
        writer.writerows(database)

# Função de Extração do Texto
def extrair_texto_do_docx(caminho_arquivo):
    texto = ""
    try:
        doc = Document(caminho_arquivo)
        for paragrafo in doc.paragraphs:
            texto += paragrafo.text + "\n"

        for tabela in doc.tables:
            for linha in tabela.rows:
                for célula in linha.cells:
                    texto += célula.text + "\n"

        for seção_cabeçalho in doc.sections:
            for cabeçalho in seção_cabeçalho.header.paragraphs:
                texto += cabeçalho.text + "\n"

        for seção_rodapé in doc.sections:
            for rodapé in seção_rodapé.footer.paragraphs:
                texto += rodapé.text + "\n"
        return texto
    except:
        return texto

# Função de Raspagem de Dados
def raspagem_texto(texto, caminho):
    texto = texto.upper()
    if texto.find('DATA:') != -1:
        coluna_data = texto.split('DATA:',1)[1][:12]
        coluna_data = re.sub(r"\s+", "-", coluna_data.strip())

    if re.sub(r"\s+", "", texto).find('GÊNERO:X') != -1:
        #coluna_genero = texto.split('Gênero: ', 1)[1][:12]
        coluna_genero = "M"
    elif re.sub(r"\s+", "", texto).find('GÊNERO:MXF') != -1:
        coluna_genero = "F"
    else:
        coluna_genero = " "

    if texto.find('IDADE:') != -1:
        coluna_idade = texto.split('IDADE:',1)[1][:3]
        coluna_idade = re.sub(r"\s+", "-", coluna_idade.strip())

    if texto.find('PAÍS:') != -1:
        coluna_pais = texto.split('PAÍS:',1)[1]
        coluna_pais = coluna_pais.split('RELIGIÃO:',1)[0]
        coluna_pais = re.sub(r"\s+", "-", coluna_pais.strip())

    if texto.find('RELIGIÃO:') != -1:
        coluna_religiao = texto.split('RELIGIÃO:', 1)[1]
        coluna_religiao = coluna_religiao.split('RELIGIÃO:', 1)[0]
        coluna_religiao = re.sub(r"\s+", "-", coluna_religiao.strip())

    if re.sub(r"\s+", "", texto).find('XBRANCA') != -1:
        coluna_cor = "BRANCA"
    elif re.sub(r"\s+", "", texto).find('XPRETA') != -1:
        coluna_cor = "PRETA"
    elif re.sub(r"\s+", "", texto).find('XPARDA') != -1:
        coluna_cor = "PARDA"
    elif re.sub(r"\s+", "", texto).find('XINDÍGENA') != -1:
        coluna_cor = "INDÍGENA"
    else:
        coluna_cor = "NI"

    if re.sub(r"\s+", "", texto).find('XNÃOPOSSUIDOCUMENTO') != -1:
        coluna_doc = "N"
    else:
        coluna_doc = "S"

    if re.sub(r"\s+", "", texto).find('FILHOS:XSIM') != -1:
        coluna_filhos = "S"
    else:
        coluna_filhos = "N"

    #**************************************************************************SAUDE
    #PCD OU ESPECIAL
    if re.sub(r"\s+", "", texto).find('CUIDADOESPECIAL:XSIM') != -1:
        coluna_saude_pcd = "S"
    elif re.sub(r"\s+", "", texto).find('CUIDADOESPECIAL:SIMXNÃO') != -1:
        coluna_saude_pcd = "N"
    else:
        coluna_saude_pcd = "NI"
    #PROB. SAUDE OU DOENÇA
    if re.sub(r"\s+", "", texto).find('DOENÇACRÔNICA?XSIM') != -1:
        coluna_saude_cronica = "S"
    elif re.sub(r"\s+", "", texto).find('DOENÇACRÔNICA?SIMXNÃO') != -1:
        coluna_saude_cronica = "N"
    else:
        coluna_saude_cronica = "NI"
    #REALIZANDO TRAT. SAÚDE
    if re.sub(r"\s+", "", texto).find('TRATAMENTODESAÚDE?XSIM') != -1:
        coluna_saude_tratamento = "S"
    elif re.sub(r"\s+", "", texto).find('TRATAMENTODESAÚDE?SIMXNÃO') != -1:
        coluna_saude_tratamento = "N"
    else:
        coluna_saude_tratamento = "NI"
    #FAZ USO DE MEDICAMENTOS
    if re.sub(r"\s+", "", texto).find('USODEMEDICAMENTOS?XSIM') != -1:
        coluna_saude_medicamentos = "S"
    elif re.sub(r"\s+", "", texto).find('USODEMEDICAMENTOS?SIMXNÃO') != -1:
        coluna_saude_medicamentos = "N"
    else:
        coluna_saude_medicamentos = "NI"
    #REALIZOU ALGUMA CIRUGIA
    if re.sub(r"\s+", "", texto).find('ALGUMACIRURGIA?XSIM') != -1:
        coluna_saude_cirurgia = "S"
    elif re.sub(r"\s+", "", texto).find('ALGUMACIRURGIA?SIMXNÃO') != -1:
        coluna_saude_cirurgia = "N"
    else:
        coluna_saude_cirurgia = "NI"
    #ACOMPANHAMENTO PSICOLOGICO
    if re.sub(r"\s+", "", texto).find('ACOMPANHAMENTOPSICOLÓGICO?XSIM') != -1:
        coluna_saude_psicologico = "S"
    elif re.sub(r"\s+", "", texto).find('ACOMPANHAMENTOPSICOLÓGICO?SIMXNÃO') != -1:
        coluna_saude_psicologico = "N"
    else:
        coluna_saude_psicologico = "NI"
    #ATENDIMENTO SAUDE MENTAL/PSIQUIATRIA?
    if re.sub(r"\s+", "", texto).find('SAÚDEMENTAL/PSIQUIATRIA?XSIM') != -1:
        coluna_saude_mental = "S"
    elif re.sub(r"\s+", "", texto).find('SAÚDEMENTAL/PSIQUIATRIA?SIMXNÃO') != -1:
        coluna_saude_mental = "N"
    else:
        coluna_saude_mental = "NI"
    #USO DE DROGA?
    if re.sub(r"\s+", "", texto).find('ALGUMTIPODEDROGA?XSIM') != -1:
        coluna_saude_droga = "S"
    elif re.sub(r"\s+", "", texto).find('ALGUMTIPODEDROGA?SIMXNÃO') != -1:
        coluna_saude_droga = "N"
    else:
        coluna_saude_droga = "NI"
    #TRATAMENTO PARA DEPENDÊNCIA QUÍMICA?
    if re.sub(r"\s+", "", texto).find('TRATAMENTOPARADEPENDÊNCIAQUÍMICA?XSIM') != -1:
        coluna_saude_quimica = "S"
    elif re.sub(r"\s+", "", texto).find('TRATAMENTOPARADEPENDÊNCIAQUÍMICA?SIMXNÃO') != -1:
        coluna_saude_quimica = "N"
    else:
        coluna_saude_quimica = "NI"


    if re.sub(r"\s+", "", texto).find('XNÃOALFABETIZADO') != -1:
        coluna_escola = "NÃO ALFABETIZADO"
    elif re.sub(r"\s+", "", texto).find('XFUNDAMENTALCOMPLETO') != -1:
        coluna_escola = "FUNDAMENTAL COMPLETO"
    elif re.sub(r"\s+", "", texto).find('XFUNDAMENTALINCOMPLETO') != -1:
        coluna_escola = "FUNDAMENTAL INCOMPLETO "
    elif re.sub(r"\s+", "", texto).find('XMÉDIOCOMPLETO') != -1:
        coluna_escola = "MÉDIO COMPLETO"
    elif re.sub(r"\s+", "", texto).find('XMÉDIOINCOMPLETO') != -1:
        coluna_escola = "MÉDIO INCOMPLETO"
    elif re.sub(r"\s+", "", texto).find('XSUPERIORCOMPLETO') != -1:
        coluna_escola = "SUPERIOR COMPLETO"
    elif re.sub(r"\s+", "", texto).find('XSUPERIORINCOMPLETO') != -1:
        coluna_escola = "SUPERIOR INCOMPLETO"
    else:
        coluna_escola = "NI"

    if re.sub(r"\s+", "", texto).find('ATIVIDADEREMUNERADA?XSIM') != -1:
        coluna_trabalho = "S"
    else:
        coluna_trabalho = "N"

    #BENEFICIOS
    #RECEBE BOLSA FAMÍLIA?
    if re.sub(r"\s+", "", texto).find('RECEBEBOLSAFAMÍLIA?XSIM') != -1:
        coluna_beneficio_bolsa = "S"
    elif re.sub(r"\s+", "", texto).find('RECEBEBOLSAFAMÍLIA?SIMXNÃO') != -1:
        coluna_beneficio_bolsa = "N"
    else:
        coluna_beneficio_bolsa = "NI"
    #TEM ID JOVEM?
    if re.sub(r"\s+", "", texto).find('TEMIDJOVEM?XSIM') != -1:
        coluna_beneficio_idjovem = "S"
    elif re.sub(r"\s+", "", texto).find('TEMIDJOVEM?SIMXNÃO') != -1:
        coluna_beneficio_idjovem = "N"
    else:
        coluna_beneficio_idjovem = "NI"
    #RECEBE BPC?
    if re.sub(r"\s+", "", texto).find('RECEBEBPC?XSIM') != -1:
        coluna_beneficio_bps = "S"
    elif re.sub(r"\s+", "", texto).find('RECEBEBPC?SIMXNÃO') != -1:
        coluna_beneficio_bps = "N"
    else:
        coluna_beneficio_bps = "NI"
    #BENEFÍCIO DA POLÍTICA DE ASSIST. SOCIAL?
    if re.sub(r"\s+", "", texto).find('POLÍTICADEASSIST.SOCIAL?XSIM') != -1:
        coluna_beneficio_ass_soc = "S"
    elif re.sub(r"\s+", "", texto).find('POLÍTICADEASSIST.SOCIAL?SIMXNÃO') != -1:
        coluna_beneficio_ass_soc = "N"
    else:
        coluna_beneficio_ass_soc = "NI"
    #BENEFÍCIO DA PREVIDÊNCIA SOCIAL?
    if re.sub(r"\s+", "", texto).find('DAPREVIDÊNCIASOCIAL?XSIM') != -1:
        coluna_beneficio_prev_soc = "S"
    elif re.sub(r"\s+", "", texto).find('DAPREVIDÊNCIASOCIAL?SIMXNÃO') != -1:
        coluna_beneficio_prev_soc = "N"
    else:
        coluna_beneficio_prev_soc = "NI"

    #MOTIVO
    #EGRESSO DO SISTEMA PENAL?
    if re.sub(r"\s+", "", texto).find('XEGRESSODOSISTEMAPENAL') != -1:
        coluna_motivo_penal = "S"
    else:
        coluna_motivo_penal = "NI"
    #VÍNCULOS FAMILIARES ROMPIDOS?
    if re.sub(r"\s+", "", texto).find('XVÍNCULOSFAMILIARESROMPIDOS') != -1:
        coluna_motivo_familiares = "S"
    else:
        coluna_motivo_familiares = "NI"
    #ABANDONO?
    if re.sub(r"\s+", "", texto).find('XABANDONO') != -1:
        coluna_motivo_abandono = "S"
    else:
        coluna_motivo_abandono = "NI"
    #FOI ROUBADO OU FURTADO?
    if re.sub(r"\s+", "", texto).find('XFOIROUBADOOUFURTADO') != -1:
        coluna_motivo_roubado = "S"
    else:
        coluna_motivo_roubado = "NI"
    #DEPEND. ÁLCOOL E DROGA?
    if re.sub(r"\s+", "", texto).find('XDEPEND.ÁLCOOLEDROGA') != -1:
        coluna_motivo_depend = "S"
    else:
        coluna_motivo_depend = "NI"
    #ESTRANGEIRO?
    if re.sub(r"\s+", "", texto).find('XESTRANGEIRO') != -1:
        coluna_motivo_extrang = "S"
    else:
        coluna_motivo_extrang = "NI"
    #DESEMPREGO?
    if re.sub(r"\s+", "", texto).find('XDESEMPREGO') != -1:
        coluna_motivo_desemp = "S"
    else:
        coluna_motivo_desemp = "NI"
    #SAÚDE MENTAL?
    if re.sub(r"\s+", "", texto).find('XSAÚDEMENTAL') != -1:
        coluna_motivo_mental = "S"
    else:
        coluna_motivo_mental = "NI"
    #TRECHEIRO?
    if re.sub(r"\s+", "", texto).find('XTRECHEIRO') != -1:
        coluna_motivo_trecheiro = "S"
    else:
        coluna_motivo_trecheiro = "NI"
    #MIGRANTE?
    if re.sub(r"\s+", "", texto).find('XMIGRANTE') != -1:
        coluna_motivo_migrante = "S"
    else:
        coluna_motivo_migrante = "NI"


    if texto.find('SITUAÇÃO DE RUA?') != -1:
        coluna_situacao_rua = texto.split('SITUAÇÃO DE RUA?', 1)[1]
        coluna_situacao_rua = coluna_situacao_rua.split('HÁ QUANTO TEMPO ESTÁ EM SITUAÇÃO DE RUA?', 1)[0]
        coluna_situacao_rua = re.sub(r"\s+", "-", coluna_situacao_rua.strip())

    if texto.find('DE RUA EM FOZ DO IGUAÇU?') != -1:
        coluna_situacao_rua_foz = texto.split('DE RUA EM FOZ DO IGUAÇU?', 1)[1]
        coluna_situacao_rua_foz = coluna_situacao_rua_foz.split('HÁ QUANTO TEMPO ESTÁ EM SITUAÇÃO DE RUA EM FOZ DO IGUAÇU?', 1)[0]
        coluna_situacao_rua_foz = re.sub(r"\s+", "-", coluna_situacao_rua_foz.strip())

    if texto.find('DA VINDA PARA FOZ?') != -1:
        coluna_motivo_foz = texto.split('DA VINDA PARA FOZ?', 1)[1]
        coluna_motivo_foz = coluna_motivo_foz.split('QUAL O MOTIVO DA VINDA PARA FOZ?', 1)[0]
        coluna_motivo_foz = re.sub(r"\s+", "-", coluna_motivo_foz.strip())

    if texto.find('PARA OUTRA CIDADE/ESTADO/PAÍS?') != -1:
        coluna_ir_para_outro = texto.split('PARA OUTRA CIDADE/ESTADO/PAÍS?', 1)[1]
        coluna_ir_para_outro = coluna_ir_para_outro.split('PRETENDE IR PARA OUTRA CIDADE/ESTADO/PAÍS?', 1)[0]
        coluna_ir_para_outro = re.sub(r"\s+", "-", coluna_ir_para_outro.strip())

    if re.sub(r"\s+", "", texto).find('MORAVAEMCASAXPRÓPRIA') != -1:
        coluna_morava = "PRÓPRIA"
    elif re.sub(r"\s+", "", texto).find('PRÓPRIAXCEDIDA') != -1:
        coluna_morava = "CEDIDA"
    elif re.sub(r"\s+", "", texto).find('CEDIDAXALUGADA') != -1:
        coluna_morava = "ALUGADA"
    else:
        coluna_morava = "NI"

    try:
        return [caminho, coluna_data, coluna_genero, coluna_idade, coluna_pais, coluna_religiao, coluna_cor,
                coluna_doc, coluna_filhos, coluna_saude_pcd, coluna_saude_cronica, coluna_saude_tratamento,
                coluna_saude_medicamentos, coluna_saude_cirurgia,
                coluna_saude_psicologico, coluna_saude_mental, coluna_saude_droga, coluna_saude_quimica,
                coluna_escola, coluna_trabalho, coluna_beneficio_bolsa, coluna_beneficio_idjovem, coluna_beneficio_bps,
                coluna_beneficio_ass_soc, coluna_beneficio_prev_soc, coluna_motivo_penal, coluna_motivo_familiares,
                coluna_motivo_abandono, coluna_motivo_roubado, coluna_motivo_depend, coluna_motivo_extrang,
                coluna_motivo_desemp,
                coluna_motivo_mental, coluna_motivo_trecheiro, coluna_motivo_migrante, coluna_situacao_rua,
                coluna_situacao_rua_foz,
                coluna_motivo_foz, coluna_ir_para_outro, coluna_morava]
    except:
        return []

#caminho_origem = r'C:\PROJETOS\KAMILA'
caminho_origem = input("Informe o Caminho de Origem: ")
percorrer_pasta(caminho_origem)

input("Script de leitura finalizado com sucesso!")
